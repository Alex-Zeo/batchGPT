# mypy: ignore-errors
import os
import re
import io
import uuid
import tempfile
import zipfile
from typing import Dict, Optional, Tuple, List, Any
from logs.logger import logger, setup_logger as _setup_logger
import concurrent.futures
import mimetypes

import pandas as pd

from .excelreader import ExcelReader

# File handling imports
import pdfplumber
import docx
from datetime import datetime


def setup_logger(log_path: str = "logs"):
    """Delegate to the package logger configuration."""
    _setup_logger(log_path)
    logger.info("File processor logger initialized")
    return logger


def extract_text_from_pdf(
    file_content, extract_tables: bool = False
) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text from PDF file
    Returns a tuple of (extracted_text, metadata)
    """
    logger.info("Extracting text from PDF")
    text = ""
    metadata = {"pages": 0, "tables": 0, "has_images": False}

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(file_content.getvalue())
            temp_file_path = temp_file.name

        # Extract text with pdfplumber
        with pdfplumber.open(temp_file_path) as pdf:
            metadata["pages"] = len(pdf.pages)

            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                text += page_text + "\n\n"

                # Check for images
                if not metadata["has_images"] and page.images:
                    metadata["has_images"] = True

                # Extract tables if requested
                if extract_tables:
                    tables = page.extract_tables()
                    if tables:
                        metadata["tables"] += len(tables)
                        for table_num, table in enumerate(tables):
                            if table:
                                table_text = (
                                    "\nTable "
                                    + str(page_num + 1)
                                    + "-"
                                    + str(table_num + 1)
                                    + ":\n"
                                )
                                for row in table:
                                    # Filter out None values and join cells
                                    row_text = " | ".join(
                                        [
                                            str(cell) if cell is not None else ""
                                            for cell in row
                                        ]
                                    )
                                    table_text += row_text + "\n"
                                text += table_text + "\n"

        # Clean up temp file
        os.unlink(temp_file_path)
        logger.info(
            f"Successfully extracted {len(text)} characters from PDF with {metadata['pages']} pages"
        )
        return text, metadata
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        if "temp_file_path" in locals():
            try:
                os.unlink(temp_file_path)
            except Exception:
                pass
        return f"Error extracting text from PDF: {str(e)}", {"error": str(e)}


def extract_text_from_docx(file_content) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text from DOCX file
    Returns a tuple of (extracted_text, metadata)
    """
    logger.info("Extracting text from DOCX")
    text = ""
    metadata = {"paragraphs": 0, "tables": 0, "has_images": False}

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            temp_file.write(file_content.getvalue())
            temp_file_path = temp_file.name

        # Extract text with python-docx
        doc = docx.Document(temp_file_path)

        # Get paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
                metadata["paragraphs"] += 1

        # Get table text
        for table in doc.tables:
            metadata["tables"] += 1
            table_text = "\nTable " + str(metadata["tables"]) + ":\n"
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                table_text += row_text + "\n"
            text += table_text + "\n"

        # Check for images (approximation)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                metadata["has_images"] = True
                break

        # Clean up temp file
        os.unlink(temp_file_path)
        logger.info(f"Successfully extracted {len(text)} characters from DOCX")
        return text, metadata
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        if "temp_file_path" in locals():
            try:
                os.unlink(temp_file_path)
            except Exception:
                pass
        return f"Error extracting text from DOCX: {str(e)}", {"error": str(e)}


def extract_text_from_txt(file_content) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text from plain text file
    Returns a tuple of (extracted_text, metadata)
    """
    logger.info("Extracting text from TXT")
    try:
        # For text files, we can directly decode the content
        text = file_content.getvalue().decode("utf-8")
        lines = text.count("\n") + 1
        metadata = {"lines": lines, "characters": len(text), "words": len(text.split())}
        logger.info(f"Successfully extracted {len(text)} characters from TXT")
        return text, metadata
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {str(e)}")
        return f"Error extracting text from TXT: {str(e)}", {"error": str(e)}


def extract_text_from_py(file_content) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text from Python file (preserving formatting)
    Returns a tuple of (extracted_text, metadata)
    """
    logger.info("Extracting text from Python file")
    try:
        # For Python files, we directly decode but preserve the formatting
        text = file_content.getvalue().decode("utf-8")
        lines = text.count("\n") + 1

        # Count classes and functions (simple approximation)
        class_count = len(re.findall(r"^\s*class\s+\w+", text, re.MULTILINE))
        function_count = len(re.findall(r"^\s*def\s+\w+", text, re.MULTILINE))

        metadata = {
            "lines": lines,
            "characters": len(text),
            "classes": class_count,
            "functions": function_count,
        }
        logger.info(f"Successfully extracted {len(text)} characters from Python file")
        return text, metadata
    except Exception as e:
        logger.error(f"Error extracting text from Python file: {str(e)}")
        return f"Error extracting text from Python file: {str(e)}", {"error": str(e)}


def extract_text_from_code(file_content, extension) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text from generic code file
    Returns a tuple of (extracted_text, metadata)
    """
    logger.info(f"Extracting text from {extension.upper()} file")
    try:
        # Decode the content preserving formatting
        text = file_content.getvalue().decode("utf-8")
        lines = text.count("\n") + 1

        metadata = {"lines": lines, "characters": len(text), "extension": extension}
        logger.info(
            f"Successfully extracted {len(text)} characters from {extension.upper()} file"
        )
        return text, metadata
    except Exception as e:
        logger.error(f"Error extracting text from {extension.upper()} file: {str(e)}")
        return f"Error extracting text from {extension.upper()} file: {str(e)}", {
            "error": str(e)
        }


def extract_data_from_excel(file_content) -> Tuple[str, Dict[str, Any]]:
    """Extract data from an Excel file using :class:`ExcelReader`."""

    logger.info("Extracting data from Excel file")
    reader = ExcelReader()
    try:
        records = reader.read(file_content)
        if records:
            df = pd.DataFrame(records)
            text = df.to_csv(index=False)
        else:
            text = ""
        metadata = {"rows": len(records), "columns": len(records[0]) if records else 0}
        logger.info(f"Successfully extracted {metadata['rows']} rows from Excel file")
        return text, metadata
    except Exception as e:  # noqa: BLE001
        logger.error(f"Error extracting data from Excel file: {str(e)}")
        return f"Error extracting data from Excel file: {str(e)}", {"error": str(e)}


def process_uploaded_file(
    file, file_type: Optional[str] = None, extract_tables: bool = False
) -> Tuple[str, str, Dict[str, Any]]:
    """
    Process uploaded file and extract text based on file type
    Returns a tuple of (extracted_text, file_type, metadata)
    """
    if file is None:
        logger.warning("No file provided")
        return "", "", {}

    # Determine file type if not provided
    if file_type is None:
        file_name = file.name.lower()

        # Use mimetypes to get better file type detection
        mime_type, _ = mimetypes.guess_type(file_name)

        if file_name.endswith(".pdf"):
            file_type = "pdf"
        elif file_name.endswith(".docx") or file_name.endswith(".doc"):
            file_type = "docx"
        elif file_name.endswith(".xlsx") or file_name.endswith(".xls"):
            file_type = "excel"
        elif file_name.endswith(".txt"):
            file_type = "txt"
        elif file_name.endswith(".py"):
            file_type = "py"
        elif file_name.endswith(
            (".js", ".html", ".css", ".json", ".ts", ".jsx", ".tsx")
        ):
            file_type = "code"
            # Store the actual extension for specific processing
            extension = os.path.splitext(file_name)[1][1:]
        else:
            # Default to text for unknown types, but track the actual extension
            file_type = "txt"
            extension = os.path.splitext(file_name)[1][1:] if "." in file_name else ""

    # Extract text based on file type
    metadata = {}
    if file_type == "pdf":
        text, metadata = extract_text_from_pdf(file, extract_tables)
    elif file_type == "docx":
        text, metadata = extract_text_from_docx(file)
    elif file_type == "excel":
        text, metadata = extract_data_from_excel(file)
    elif file_type == "py":
        text, metadata = extract_text_from_py(file)
    elif file_type == "code":
        text, metadata = extract_text_from_code(file, extension)
    else:  # Default to txt
        text, metadata = extract_text_from_txt(file)

    # Add filename and size to metadata
    metadata["filename"] = file.name
    metadata["file_size"] = len(file.getvalue())
    metadata["processed_at"] = datetime.now().isoformat()

    return text, file_type, metadata


def process_multiple_files(
    files, extract_tables: bool = False, process_mode: str = "separate"
) -> List[Dict[str, Any]]:
    """
    Process multiple uploaded files in parallel

    Args:
        files: List of file objects
        extract_tables: Whether to extract tables from PDFs and DOCX files
        process_mode: How to process files - "separate" (each file as separate prompt)
                      or "combine" (all files combined) or "zip" (handle zip archives)

    Returns:
        List of dictionaries with processed file information
    """
    logger.info(f"Processing {len(files)} files with mode: {process_mode}")
    results = []

    # Handle zip files if mode is "zip"
    if process_mode == "zip" and len(files) > 0:
        zip_results = []
        for file in files:
            if file.name.lower().endswith(".zip"):
                # Process zip file
                zip_results.extend(process_zip_file(file))
            else:
                # Process non-zip file normally
                text, file_type, metadata = process_uploaded_file(
                    file, extract_tables=extract_tables
                )
                zip_results.append(
                    {
                        "filename": file.name,
                        "text": text,
                        "file_type": file_type,
                        "metadata": metadata,
                        "id": str(uuid.uuid4()),
                    }
                )
        return zip_results

    # Process files in parallel using ThreadPoolExecutor
    with concurrent.futures.ThreadPoolExecutor() as executor:
        # Start the processing tasks
        future_to_file = {
            executor.submit(process_uploaded_file, file, None, extract_tables): file
            for file in files
        }

        # Collect results as they complete
        for future in concurrent.futures.as_completed(future_to_file):
            file = future_to_file[future]
            try:
                text, file_type, metadata = future.result()
                results.append(
                    {
                        "filename": file.name,
                        "text": text,
                        "file_type": file_type,
                        "metadata": metadata,
                        "id": str(uuid.uuid4()),
                    }
                )
            except Exception as e:
                logger.error(f"Error processing file {file.name}: {str(e)}")
                results.append(
                    {
                        "filename": file.name,
                        "text": f"Error processing file: {str(e)}",
                        "file_type": "error",
                        "metadata": {"error": str(e)},
                        "id": str(uuid.uuid4()),
                    }
                )

    # If we want to combine all files into one result
    if process_mode == "combine" and results:
        combined_text = (
            "\n\n==== FILE: {} ====\n\n".format(results[0]["filename"])
            + results[0]["text"]
        )
        combined_metadata = {
            "files": [r["filename"] for r in results],
            "file_types": [r["file_type"] for r in results],
            "total_files": len(results),
        }

        # Add text from remaining files
        for result in results[1:]:
            combined_text += (
                f"\n\n==== FILE: {result['filename']} ====\n\n{result['text']}"
            )

        return [
            {
                "filename": "combined_files.txt",
                "text": combined_text,
                "file_type": "combined",
                "metadata": combined_metadata,
                "id": str(uuid.uuid4()),
            }
        ]

    return results


def process_zip_file(zip_file) -> List[Dict[str, Any]]:
    """
    Extract and process files from a zip archive
    Returns a list of processed file results
    """
    logger.info(f"Processing ZIP file: {zip_file.name}")
    results = []

    try:
        # Create a temporary directory to extract files
        with tempfile.TemporaryDirectory() as temp_dir:
            # Save the zip file locally
            zip_path = os.path.join(temp_dir, "archive.zip")
            with open(zip_path, "wb") as f:
                f.write(zip_file.getvalue())

            # Extract all files
            with zipfile.ZipFile(zip_path, "r") as zip_ref:
                zip_ref.extractall(temp_dir)

            # Process each extracted file
            extracted_files = []
            for root, _, files in os.walk(temp_dir):
                for filename in files:
                    # Skip the archive.zip itself
                    if filename == "archive.zip":
                        continue

                    file_path = os.path.join(root, filename)
                    # Get relative path for cleaner display
                    rel_path = os.path.relpath(file_path, temp_dir)

                    # Skip files that are too large or binary
                    file_size = os.path.getsize(file_path)
                    if file_size > 10 * 1024 * 1024:  # Skip files larger than 10MB
                        logger.warning(
                            f"Skipping large file in ZIP: {rel_path} ({file_size} bytes)"
                        )
                        continue

                    # Try to determine if it's a text file
                    mime_type, _ = mimetypes.guess_type(filename)
                    is_text = False

                    # Check mime type or extensions for text files
                    if mime_type and mime_type.startswith("text/"):
                        is_text = True
                    elif any(
                        filename.lower().endswith(ext)
                        for ext in (
                            ".txt",
                            ".py",
                            ".js",
                            ".html",
                            ".css",
                            ".json",
                            ".md",
                            ".xml",
                            ".csv",
                            ".docx",
                            ".pdf",
                            ".c",
                            ".cpp",
                            ".java",
                        )
                    ):
                        is_text = True

                    if is_text:
                        extracted_files.append((rel_path, file_path))

            # Process files that we think are text-based
            for rel_path, file_path in extracted_files:
                try:
                    # Determine file type
                    ext = os.path.splitext(rel_path)[1].lower()
                    file_type = None
                    if ext == ".pdf":
                        file_type = "pdf"
                    elif ext in (".docx", ".doc"):
                        file_type = "docx"
                    elif ext == ".py":
                        file_type = "py"
                    elif ext in (
                        ".js",
                        ".html",
                        ".css",
                        ".json",
                        ".ts",
                        ".jsx",
                        ".tsx",
                    ):
                        file_type = "code"
                    else:
                        file_type = "txt"

                    # Open and process the file
                    with open(file_path, "rb") as f:
                        file_content = io.BytesIO(f.read())
                        file_content.name = os.path.basename(file_path)

                        # Create a custom class to mimic a file upload object
                        class FileObj:
                            def __init__(self, content, name):
                                self.content = content
                                self.name = name

                            def getvalue(self):
                                return self.content.getvalue()

                        file_obj = FileObj(file_content, file_content.name)

                        # Process the file
                        text, detected_type, metadata = process_uploaded_file(
                            file_obj, file_type
                        )

                        results.append(
                            {
                                "filename": rel_path,  # Use relative path within zip
                                "text": text,
                                "file_type": detected_type,
                                "metadata": metadata,
                                "id": str(uuid.uuid4()),
                                "from_zip": True,
                            }
                        )
                except Exception as e:
                    logger.error(f"Error processing file {rel_path} from ZIP: {str(e)}")
                    results.append(
                        {
                            "filename": rel_path,
                            "text": f"Error processing file from ZIP: {str(e)}",
                            "file_type": "error",
                            "metadata": {"error": str(e)},
                            "id": str(uuid.uuid4()),
                            "from_zip": True,
                        }
                    )

    except Exception as e:
        logger.error(f"Error processing ZIP file {zip_file.name}: {str(e)}")
        results.append(
            {
                "filename": zip_file.name,
                "text": f"Error processing ZIP file: {str(e)}",
                "file_type": "error",
                "metadata": {"error": str(e)},
                "id": str(uuid.uuid4()),
            }
        )

    return results


def split_text_into_chunks(
    text: str, max_chunk_size: int = 4000, overlap: int = 200
) -> List[str]:
    """
    Split text into chunks for processing with a specified overlap.
    Tries to split on paragraph boundaries when possible.
    """
    logger.info(
        f"Splitting text into chunks (max size: {max_chunk_size}, overlap: {overlap})"
    )

    # If text is shorter than max_chunk_size, return it as is
    if len(text) <= max_chunk_size:
        return [text]

    chunks = []
    # Split text into paragraphs (respecting newlines)
    paragraphs = text.split("\n\n")

    current_chunk = ""
    for paragraph in paragraphs:
        # If adding this paragraph would exceed max_chunk_size
        if len(current_chunk) + len(paragraph) > max_chunk_size:
            # If current_chunk is not empty, add it to chunks
            if current_chunk:
                chunks.append(current_chunk)

            # If this paragraph itself is too long, split it by sentence
            if len(paragraph) > max_chunk_size:
                sentences = re.split(r"(?<=[.!?])\s+", paragraph)
                current_chunk = ""

                for sentence in sentences:
                    # If adding this sentence would exceed max_chunk_size
                    if len(current_chunk) + len(sentence) > max_chunk_size:
                        # If current_chunk is not empty, add it to chunks
                        if current_chunk:
                            chunks.append(current_chunk)

                        # If this sentence itself is too long, split it by words
                        if len(sentence) > max_chunk_size:
                            words = sentence.split()
                            current_chunk = ""

                            for word in words:
                                if len(current_chunk) + len(word) > max_chunk_size:
                                    chunks.append(current_chunk)
                                    current_chunk = word + " "
                                else:
                                    current_chunk += word + " "
                        else:
                            current_chunk = sentence + " "
                    else:
                        current_chunk += sentence + " "
            else:
                current_chunk = paragraph
        else:
            # If adding this paragraph wouldn't exceed max_chunk_size
            if current_chunk:
                current_chunk += "\n\n" + paragraph
            else:
                current_chunk = paragraph

    # Add the last chunk if it's not empty
    if current_chunk:
        chunks.append(current_chunk)

    # Add overlap between chunks if there are multiple chunks
    if len(chunks) > 1 and overlap > 0:
        overlapped_chunks = []
        for i, chunk in enumerate(chunks):
            if i > 0:
                # Add overlap from previous chunk
                prev_chunk = chunks[i - 1]
                overlap_text = (
                    prev_chunk[-overlap:] if len(prev_chunk) > overlap else prev_chunk
                )
                chunk = overlap_text + chunk

            overlapped_chunks.append(chunk)
        chunks = overlapped_chunks

    logger.info(f"Split text into {len(chunks)} chunks")
    return chunks


def detect_file_type(filename: str) -> str:
    """Detect file type from filename"""
    extension = filename.lower().split(".")[-1] if "." in filename else ""

    if extension == "pdf":
        return "pdf"
    elif extension in ("doc", "docx"):
        return "docx"
    elif extension in ("xls", "xlsx"):
        return "excel"
    elif extension == "py":
        return "py"
    elif extension in ("js", "html", "css", "json", "ts", "jsx", "tsx"):
        return "code"
    elif extension == "zip":
        return "zip"
    else:
        return "txt"


def generate_summary(processed_files: List[Dict]) -> str:
    """Generate a summary of processed files"""
    if not processed_files:
        return "No files processed."

    summary = f"Processed {len(processed_files)} files:\n\n"

    for i, file in enumerate(processed_files, 1):
        filename = file.get("filename", "Unknown")
        file_type = file.get("file_type", "Unknown")
        metadata = file.get("metadata", {})

        summary += f"{i}. {filename} ({file_type.upper()})\n"

        # Include relevant metadata based on file type
        if file_type == "pdf":
            pages = metadata.get("pages", "Unknown")
            tables = metadata.get("tables", 0)
            has_images = metadata.get("has_images", False)
            summary += (
                f"   Pages: {pages}, Tables: {tables}, Contains Images: {has_images}\n"
            )
        elif file_type == "docx":
            paragraphs = metadata.get("paragraphs", "Unknown")
            tables = metadata.get("tables", 0)
            summary += f"   Paragraphs: {paragraphs}, Tables: {tables}\n"
        elif file_type in ("txt", "py", "code"):
            lines = metadata.get("lines", "Unknown")
            chars = metadata.get("characters", "Unknown")
            summary += f"   Lines: {lines}, Characters: {chars}\n"

        # Show text size
        text_size = len(file.get("text", ""))
        summary += f"   Extracted: {text_size} characters\n"

        if i < len(processed_files):
            summary += "\n"

    return summary


# Set up the logger when this module is imported
setup_logger()
